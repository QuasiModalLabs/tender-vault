"""
Tests for manually-dropped tender documents: manifest, dispatch table, reads.

Runs with plain Python — no pytest needed:

    python tests/test_attachments.py

Exit code 0 = all passed. Any assertion failure = non-zero + traceback.

NO REAL PDF ANYWHERE IN HERE, and that is deliberate three times over. CI would
otherwise need pdfplumber installed to test anything; a binary fixture would
have to be committed; and we cannot commit a third-party tender document in the
first place — that is the whole premise of the feature.

So the suite registers its own extractors into `attachments.EXTRACTORS` and
exercises the machinery through those. That is not a workaround: the claim
under test is that adding a format is one function plus one row, and the only
honest way to test that claim is to add one.
"""
from __future__ import annotations

import sys
from pathlib import Path
from types import SimpleNamespace

sys.path.insert(0, str(Path(__file__).parent))
import conftest  # noqa: E402  — MUST come first: redirects the vault on import
import attachments as at  # noqa: E402
import tender_tools as tt  # noqa: E402


TENDER_ID = "TEST-ATT-0001"
NOTE_STEM = "test-att-0001"          # what _slugify(TENDER_ID) produces

NOTE = f"""---
tender_id: {TENDER_ID}
title: "Attachment fixture tender"
status: watching
---

# Attachment fixture tender

## My notes
"""


def _register_text_extractor(ext: str, version: str = "1.0") -> None:
    """
    The dispatch-table claim under test: one function, one row.

    Page count is faked as 1 so the no_text_layer threshold is a flat 50
    non-whitespace characters, which makes the boundary easy to write against.
    """
    def _extract(path: Path) -> at.ExtractResult:
        return at.ExtractResult(
            text=path.read_text(encoding="utf-8"),
            page_count=1,
            extractor=f"faketext {version}",
        )

    at.EXTRACTORS[ext] = _extract


def setup_temp_vault() -> Path:
    """A vault with one watching note. The redirect itself lives in conftest."""
    root = conftest.redirect_vault()
    tt.paths.WATCHING.mkdir(parents=True)
    (tt.paths.WATCHING / f"{NOTE_STEM}.md").write_text(NOTE, encoding="utf-8", newline="\n")

    _register_text_extractor(".txt")
    # .later is deliberately NOT registered yet — test_extractor_added_later
    # registers it to prove an already-dropped file is picked up with no re-drop.
    at.EXTRACTORS.pop(".later", None)
    return root


def _folder() -> Path:
    return tt.paths.WATCHING / NOTE_STEM


def _drop(name: str, content: str) -> Path:
    path = _folder() / name
    path.write_text(content, encoding="utf-8", newline="\n")
    return path


def _list() -> dict:
    return tt.cmd_list_attachments(SimpleNamespace(tender_id=TENDER_ID))


def _record(result: dict, filename: str) -> dict:
    for rec in result["files"]:
        if rec["filename"] == filename:
            return rec
    raise AssertionError(f"{filename} not in manifest: {result['files']}")


# ---------------------------------------------------------------------------


def test_attach_creates_folder():
    result = tt.cmd_attach(SimpleNamespace(
        tender_id=TENDER_ID, platform="merx", no_reveal=True,
    ))
    assert "error" not in result, f"attach failed: {result}"
    assert result["created"] is True, "attach did not report creating the folder"

    folder = _folder()
    assert folder.is_dir(), "attachment folder not created"
    assert (folder / at.EXTRACTED_DIRNAME).is_dir(), ".extracted/ not created"
    assert (folder / at.MANIFEST_NAME).exists(), "_index.md not written"

    # The contract with the user: a path they can paste somewhere.
    assert Path(result["attachment_folder"]).is_absolute(), (
        f"attach must return an absolute path, got {result['attachment_folder']}"
    )
    assert result["source_platform"] == "merx", "provenance not recorded"


def test_attach_is_idempotent_and_keeps_provenance():
    """A second attach must not blow away a manifest that has files in it."""
    result = tt.cmd_attach(SimpleNamespace(
        tender_id=TENDER_ID, platform="ariba", no_reveal=True,
    ))
    assert result["created"] is False, "second attach reported creating the folder"
    assert result["source_platform"] == "merx", (
        "second attach overwrote the recorded source platform"
    )


def test_attach_unknown_tender_rejected():
    result = tt.cmd_attach(SimpleNamespace(
        tender_id="NO-SUCH-TENDER", platform="merx", no_reveal=True,
    ))
    assert "error" in result, "attach on a nonexistent tender should error"


def test_manifest_written_with_lf():
    raw = (_folder() / at.MANIFEST_NAME).read_bytes()
    assert b"\r\n" not in raw, "_index.md written with CRLF"


def test_extracts_a_dropped_file():
    _drop("sow.txt", "\n".join(f"Statement of work line {i}" for i in range(100)))
    result = _list()

    assert result["added"] == ["sow.txt"], f"new file not reported: {result['added']}"
    rec = _record(result, "sow.txt")
    assert rec["extraction_status"] == at.STATUS_EXTRACTED, rec
    assert rec["extractor"] == "faketext 1.0", rec
    assert rec["page_count"] == 1, rec
    assert len(rec["sha256"]) == 64, "sha256 not a full hex digest"
    assert rec["size_bytes"] > 0, rec

    text = _folder() / rec["extracted_path"]
    assert text.exists(), "no extracted text written"
    assert "Statement of work line 0" in text.read_text(encoding="utf-8")


def test_unsupported_type_still_listed():
    """
    A dropped file nothing can read must still be visible.

    The extension here has to be one with NO registered extractor. This test
    used .xlsx until a spreadsheet reader was added, at which point it kept
    passing for the wrong reason — openpyxl was failing to parse a fake
    workbook, so the assertion was measuring a corrupt file rather than an
    unsupported format. .dwg is a real thing to find in a construction annex
    and there is no plausible reader for it here.
    """
    assert ".dwg" not in at.EXTRACTORS, "pick an extension with no extractor"
    _drop("site-plan.dwg", "binary-ish CAD payload nothing here reads")
    result = _list()

    rec = _record(result, "site-plan.dwg")
    assert rec["extraction_status"] == at.STATUS_UNSUPPORTED, rec
    assert rec["extractor"] is None, rec
    assert rec["extracted_path"] is None, "unsupported type must produce no text"
    # No extractor ran, so there is nothing to warn about — this is a fact
    # about the format, not a failure.
    assert not any("site-plan.dwg" in w for w in result["warnings"]), (
        f"an unsupported format is not an error: {result['warnings']}"
    )
    # The point of listing it at all: identity and size are still recorded.
    assert len(rec["sha256"]) == 64, rec
    assert rec["size_bytes"] > 0, rec
    assert rec["extension"] == ".dwg", rec


def test_corrupt_file_is_distinguished_from_unsupported():
    """
    A file whose extractor ran and threw is not the same as one nothing reads.

    Both land on unsupported_type — the enum has three values and this is the
    honest bucket — but the corrupt one carries a warning naming the failure,
    which is the only thing telling the user to re-download it.
    """
    _drop("Broken.xlsx", "this is not a zip and openpyxl will say so")
    result = _list()

    rec = _record(result, "Broken.xlsx")
    assert rec["extraction_status"] == at.STATUS_UNSUPPORTED, rec
    assert any("Broken.xlsx" in w for w in result["warnings"]), (
        f"a file that failed to parse must say so: {result['warnings']}"
    )


def test_no_text_layer_writes_no_file():
    """
    A scan yields nothing. Recording that is the job; writing a zero-byte .txt
    would look like a successful extraction of an empty document.
    """
    _drop("scanned.txt", "   \n  \n tiny \n")
    result = _list()

    rec = _record(result, "scanned.txt")
    assert rec["extraction_status"] == at.STATUS_NO_TEXT_LAYER, rec
    assert rec["extracted_path"] is None, rec
    assert rec["char_count"] < 50, f"fixture is not below the threshold: {rec}"
    assert not (_folder() / at.EXTRACTED_DIRNAME / "scanned.txt").exists(), (
        "a zero-byte .txt was written for a document with no text layer"
    )
    assert "scan" in (rec.get("status_note") or "").lower(), (
        "no_text_layer should say what the condition is"
    )


def test_unchanged_file_is_not_re_extracted():
    before = (_folder() / at.EXTRACTED_DIRNAME / "sow.txt").stat().st_mtime_ns
    result = _list()
    assert result["added"] == [] and result["changed"] == [], (
        f"a no-op refresh reported changes: {result}"
    )
    after = (_folder() / at.EXTRACTED_DIRNAME / "sow.txt").stat().st_mtime_ns
    assert before == after, "unchanged file was extracted again"


def test_amended_document_detected_and_re_extracted():
    """
    The case the hashing exists for: MERX posts an addendum mid-solicitation
    under the same filename. Stale text that looks current is the failure.
    """
    old_sha = _record(_list(), "sow.txt")["sha256"]

    _drop("sow.txt", "\n".join(f"AMENDED line {i}" for i in range(100)))
    result = _list()

    assert result["changed"] == ["sow.txt"], (
        f"amended document not reported as changed: {result['changed']}"
    )
    rec = _record(result, "sow.txt")
    assert rec["sha256"] != old_sha, "hash did not change"
    assert rec["superseded_sha256"] == old_sha, (
        "the superseded hash is what makes an amendment reconstructable later"
    )
    text = (_folder() / rec["extracted_path"]).read_text(encoding="utf-8")
    assert "AMENDED line 0" in text, "text was not re-extracted"
    assert "Statement of work" not in text, "stale text survived the amendment"


def test_extractor_added_later_needs_no_re_drop():
    """
    THE DISPATCH-TABLE REQUIREMENT. A file dropped before its format was
    supported must be picked up once an extractor is registered — one function
    and one row, no manifest migration, and the user does not re-drop anything.
    """
    _drop("annex.later", "\n".join(f"Annex clause {i}" for i in range(40)))

    first = _record(_list(), "annex.later")
    assert first["extraction_status"] == at.STATUS_UNSUPPORTED, first
    sha_before = first["sha256"]

    _register_text_extractor(".later", version="0.1")   # the whole registration

    second = _record(_list(), "annex.later")
    assert second["extraction_status"] == at.STATUS_EXTRACTED, (
        f"registering an extractor did not pick up the existing file: {second}"
    )
    assert second["sha256"] == sha_before, "the file was not touched, only read"
    assert second["extractor"] == "faketext 0.1", second
    assert (_folder() / second["extracted_path"]).exists(), "no text written"


def test_deleted_extracted_text_is_rebuilt():
    text_path = _folder() / _record(_list(), "sow.txt")["extracted_path"]
    text_path.unlink()
    rec = _record(_list(), "sow.txt")
    assert rec["extraction_status"] == at.STATUS_EXTRACTED, rec
    assert (_folder() / rec["extracted_path"]).exists(), (
        "extracted text deleted by hand was not rebuilt"
    )


def test_removed_file_drops_out_of_manifest():
    _drop("temporary.txt", "x" * 200)
    assert _record(_list(), "temporary.txt")

    (_folder() / "temporary.txt").unlink()
    result = _list()
    assert "temporary.txt" in result["removed"], f"removal not reported: {result}"
    assert all(r["filename"] != "temporary.txt" for r in result["files"]), (
        "a deleted file is still in the manifest"
    )


def test_filename_with_awkward_characters_round_trips():
    """
    Filenames are KEYS. The manifest quotes them as JSON strings precisely so
    spaces and punctuation survive the write/read cycle exactly.
    """
    name = "RFP W2187-SPO, Annex A & B (rev 2).txt"
    _drop(name, "\n".join(f"clause {i}" for i in range(30)))
    rec = _record(_list(), name)
    assert rec["filename"] == name, f"filename did not round-trip: {rec['filename']}"
    assert rec["extraction_status"] == at.STATUS_EXTRACTED, rec


def test_read_is_paginated():
    result = tt.cmd_read_attachment(SimpleNamespace(
        tender_id=TENDER_ID, filename="sow.txt", offset=0, limit=10,
    ))
    assert "error" not in result, result
    assert result["lines_returned"] == 10, result
    assert result["total_lines"] == 100, result
    assert result["eof"] is False, "a 10-line window of 100 lines is not eof"
    assert result["text"].startswith("AMENDED line 0"), result["text"][:60]
    assert "AMENDED line 10" not in result["text"], "window overran its limit"


def test_read_offset_and_eof():
    result = tt.cmd_read_attachment(SimpleNamespace(
        tender_id=TENDER_ID, filename="sow.txt", offset=90, limit=400,
    ))
    assert result["lines_returned"] == 10, result
    assert result["eof"] is True, "reading to the end did not report eof"
    assert result["text"].startswith("AMENDED line 90"), result["text"][:60]


def test_read_limit_is_clamped():
    result = tt.cmd_read_attachment(SimpleNamespace(
        tender_id=TENDER_ID, filename="sow.txt", offset=0, limit=999999,
    ))
    assert result["limit"] == 2000, f"limit not clamped: {result['limit']}"


def test_read_rejects_a_file_with_no_text():
    result = tt.cmd_read_attachment(SimpleNamespace(
        tender_id=TENDER_ID, filename="scanned.txt", offset=0, limit=10,
    ))
    assert "error" in result, "reading a no_text_layer file should error"
    assert result["extraction_status"] == at.STATUS_NO_TEXT_LAYER, result
    assert "text" not in result, "an unreadable file must not return empty text"


def test_read_picks_up_a_change_without_a_list_call():
    """
    A read re-hashes the one file it serves, so it can never hand back text
    that no longer matches the document on disk — even if nobody called
    list-attachments in between.
    """
    _drop("sow.txt", "\n".join(f"SECOND AMENDMENT line {i}" for i in range(20)))
    result = tt.cmd_read_attachment(SimpleNamespace(
        tender_id=TENDER_ID, filename="sow.txt", offset=0, limit=5,
    ))
    assert "error" not in result, result
    assert result["text"].startswith("SECOND AMENDMENT line 0"), result["text"][:60]
    assert result["total_lines"] == 20, result


def test_same_stem_different_extension_keeps_separate_text():
    """
    REGRESSION. Two documents whose names differ only by extension must not
    share one .txt.

    The subtle half is that the collision only appeared ACROSS runs: a file
    skipped as unchanged never asks for an extracted name, so its .txt looked
    unclaimed and the next document dropped with the same stem was handed it.
    The result was two manifest records pointing at one file holding one
    document's text — a read returning the wrong document while reporting
    extraction_status: extracted, which is exactly the kind of confidently
    wrong answer the rest of this module is built to avoid.
    """
    _register_text_extractor(".docm")
    try:
        _drop("Annex-Z.txt", "ORIGINAL TXT CONTENT " * 8)
        first = _record(_list(), "Annex-Z.txt")["extracted_path"]

        # Second run, second document, same stem. The first is untouched.
        _drop("Annex-Z.docm", "DIFFERENT DOCM CONTENT " * 8)
        result = _list()
        txt = _record(result, "Annex-Z.txt")
        docm = _record(result, "Annex-Z.docm")

        assert txt["extracted_path"] != docm["extracted_path"], (
            f"two documents share one extracted file: {txt['extracted_path']}"
        )
        assert txt["extracted_path"] == first, (
            "an untouched document's extracted path moved under it"
        )
        assert "ORIGINAL TXT" in (_folder() / txt["extracted_path"]).read_text(
            encoding="utf-8"), "the first document's text was overwritten"
        assert "DIFFERENT DOCM" in (_folder() / docm["extracted_path"]).read_text(
            encoding="utf-8"), "the second document's text is wrong"
    finally:
        at.EXTRACTORS.pop(".docm", None)


def test_extracted_path_is_stable_across_amendments():
    """An amendment must reuse the same .txt, not orphan the old one."""
    before = _record(_list(), "Annex-Z.txt")["extracted_path"]
    _drop("Annex-Z.txt", "AMENDED TXT CONTENT " * 8)
    after = _record(_list(), "Annex-Z.txt")["extracted_path"]
    assert before == after, f"extracted path churned: {before} -> {after}"
    assert "AMENDED TXT" in (_folder() / after).read_text(encoding="utf-8")


def test_unknown_platform_rejected():
    """
    argparse enforces --platform choices on the CLI, but the MCP tool builds
    its own namespace and bypasses the parser entirely. Provenance that nobody
    validated is provenance nobody should trust.
    """
    result = tt.cmd_attach(SimpleNamespace(
        tender_id=TENDER_ID, platform="sharepoint", no_reveal=True,
    ))
    assert "error" in result, "an unrecognised platform should be rejected"
    assert "merx" in result["error"], "the error should name the valid options"


def test_missing_library_is_reported_not_fatal():
    """
    A registered format whose library isn't installed must not take down the
    whole listing — every other file still gets read, and the miss is a
    warning rather than a silent fact about the document.

    It also has to self-heal: the file goes to unsupported_type, which is the
    state the refresh rule already retries once an extractor works.
    """
    def _broken(path: Path) -> at.ExtractResult:
        raise ImportError("No module named 'nothinghere'")

    at.EXTRACTORS[".needslib"] = _broken
    try:
        _drop("annex-d.needslib", "content nobody can read yet " * 5)
        result = _list()

        rec = _record(result, "annex-d.needslib")
        assert rec["extraction_status"] == at.STATUS_UNSUPPORTED, rec
        assert rec["extracted_path"] is None, rec
        assert any("annex-d.needslib" in w for w in result["warnings"]), (
            f"a missing library must be surfaced as a warning: {result['warnings']}"
        )
        # The rest of the folder is unaffected.
        assert _record(result, "sow.txt")["extraction_status"] == at.STATUS_EXTRACTED

        at.EXTRACTORS[".needslib"] = lambda p: at.ExtractResult(
            text=p.read_text(encoding="utf-8"), page_count=None, extractor="fixed 1.0",
        )
        healed = _record(_list(), "annex-d.needslib")
        assert healed["extraction_status"] == at.STATUS_EXTRACTED, (
            f"installing the library did not pick the file back up: {healed}"
        )
    finally:
        at.EXTRACTORS.pop(".needslib", None)


def test_real_docx_extraction():
    """
    The one test here that uses a real library and a real document.

    Skips cleanly when python-docx isn't installed, matching how the suites
    that need a built database behave. Everything else in this file runs on
    fake extractors on purpose; this exists because the .docx extractor makes
    two claims a fake cannot check — that a table survives as rows, and that
    prose and tables come back in document order rather than in two batches.
    """
    try:
        import docx
    except ImportError:
        print("  SKIP test_real_docx_extraction (python-docx not installed)")
        return

    document = docx.Document()
    document.add_paragraph("1. Statement of work. The contractor shall deliver.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "M1"
    table.cell(0, 1).text = "Bilingual documentation"
    table.cell(1, 0).text = "M2"
    table.cell(1, 1).text = "Protected A hosting in Canada"
    document.add_paragraph("2. Evaluation. Mandatory criteria above are pass/fail.")
    document.save(str(_folder() / "SOW.docx"))

    rec = _record(_list(), "SOW.docx")
    assert rec["extraction_status"] == at.STATUS_EXTRACTED, rec
    assert rec["extractor"].startswith("python-docx "), rec
    # A .docx has no pages until something renders it. Inventing a number here
    # is exactly what the nullable field exists to avoid.
    assert rec["page_count"] is None, f"page_count should be null for .docx: {rec}"

    text = (_folder() / rec["extracted_path"]).read_text(encoding="utf-8")
    lines = [ln for ln in text.splitlines() if ln.strip()]

    assert "M1 | Bilingual documentation" in lines, (
        f"table row was not kept row-shaped: {lines}"
    )
    assert "M2 | Protected A hosting in Canada" in lines, lines
    # Document order: the table sits BETWEEN the two clauses, not after both.
    sow = next(i for i, ln in enumerate(lines) if ln.startswith("1. Statement"))
    row = next(i for i, ln in enumerate(lines) if ln.startswith("M1 |"))
    evaluation = next(i for i, ln in enumerate(lines) if ln.startswith("2. Evaluation"))
    assert sow < row < evaluation, (
        f"prose and tables came back out of document order: {lines}"
    )


def test_docx_dropped_before_support_existed():
    """
    THE UPGRADE PATH, with the real extractor.

    A .docx sitting in the folder from before .docx was registered must be
    picked up on the next call. The user does not re-drop it and the manifest
    is not migrated — the file's own hash is unchanged throughout.
    """
    try:
        import docx
    except ImportError:
        print("  SKIP test_docx_dropped_before_support_existed (no python-docx)")
        return

    document = docx.Document()
    document.add_paragraph("Annex C. " + "Deliverables and acceptance. " * 6)
    document.save(str(_folder() / "Annex-C.docx"))

    real = at.EXTRACTORS.pop(".docx")          # pretend v1: no .docx support
    try:
        before = _record(_list(), "Annex-C.docx")
        assert before["extraction_status"] == at.STATUS_UNSUPPORTED, before
        assert before["extracted_path"] is None, before
    finally:
        at.EXTRACTORS[".docx"] = real          # the entire "upgrade"

    after = _record(_list(), "Annex-C.docx")
    assert after["extraction_status"] == at.STATUS_EXTRACTED, (
        f"registering .docx did not pick up the already-dropped file: {after}"
    )
    assert after["sha256"] == before["sha256"], "the file itself was touched"
    assert "Deliverables and acceptance" in (
        (_folder() / after["extracted_path"]).read_text(encoding="utf-8")
    )


def test_real_xlsx_keeps_every_value_addressable():
    """
    The whole reason .xlsx was held back: a figure must stay attached to its
    row. This asserts the addressing survives, not merely that text came out.
    """
    try:
        import openpyxl
    except ImportError:
        print("  SKIP test_real_xlsx_keeps_every_value_addressable (no openpyxl)")
        return

    book = openpyxl.Workbook()
    sheet = book.active
    sheet.title = "Pricing"
    for row in [
        ("Ref", "Resource", "Days", "Rate"),
        ("L1", "Senior developer", 100, 950),
        ("L2", "Business analyst", 60, 720),
    ]:
        sheet.append(row)
    book.save(str(_folder() / "Pricing-Table.xlsx"))

    rec = _record(_list(), "Pricing-Table.xlsx")
    assert rec["extraction_status"] == at.STATUS_EXTRACTED, rec
    assert rec["extractor"].startswith("openpyxl "), rec
    assert rec["page_count"] is None, f"sheets are not pages: {rec}"

    lines = (_folder() / rec["extracted_path"]).read_text(
        encoding="utf-8").splitlines()

    assert any("Pricing" in ln and "Sheet 1 of 1" in ln for ln in lines), (
        f"sheet identity missing: {lines[:4]}"
    )
    assert any(ln.startswith("cols | A | B | C | D") for ln in lines), (
        f"column letters missing, so no value can be addressed: {lines[:6]}"
    )
    # The claim that matters: rate 950 is on the row that says "Senior
    # developer", and it is row 2, not "somewhere in the document".
    senior = next(ln for ln in lines if "Senior developer" in ln)
    assert senior.startswith("r2 | "), f"row number lost: {senior!r}"
    assert senior.endswith("| 950"), f"rate detached from its row: {senior!r}"
    analyst = next(ln for ln in lines if "Business analyst" in ln)
    assert analyst.startswith("r3 | ") and analyst.endswith("| 720"), analyst


def test_xlsx_uncalculated_formula_is_labelled_not_blank():
    """
    THE TRAP. openpyxl returns the value Excel cached on its last save. A
    workbook written by a script has no cache, so every formula reads as None
    — and a pricing column of blanks looks like the bidder left it empty
    rather than like a number we failed to read.
    """
    try:
        import openpyxl
    except ImportError:
        print("  SKIP test_xlsx_uncalculated_formula_is_labelled_not_blank")
        return

    book = openpyxl.Workbook()
    sheet = book.active
    sheet.title = "Totals"
    sheet.append(("Item", "Qty", "Unit", "Total"))
    sheet.append(("Licences", 10, 500, "=B2*C2"))     # never calculated
    book.save(str(_folder() / "Totals.xlsx"))

    rec = _record(_list(), "Totals.xlsx")
    text = (_folder() / rec["extracted_path"]).read_text(encoding="utf-8")

    assert "<uncalculated formula>" in text, (
        f"an uncalculated formula rendered as blank: {text!r}"
    )
    # The banner has to name the cell and say what the blank is NOT.
    assert "Totals!D2" in text, f"the trap is not located: {text!r}"
    assert "NOT zero" in text, f"the banner does not say what blank isn't: {text!r}"
    # And it must lead, not trail.
    assert text.splitlines()[0].startswith("!!"), (
        "the warning must be read before the numbers are"
    )


def test_xlsx_hidden_sheet_is_marked():
    """A hidden pricing tab read as though it were live is its own kind of wrong."""
    try:
        import openpyxl
    except ImportError:
        print("  SKIP test_xlsx_hidden_sheet_is_marked")
        return

    book = openpyxl.Workbook()
    book.active.title = "Visible"
    book.active.append(("live", "data", "here"))
    hidden = book.create_sheet("Draft-Pricing")
    hidden.append(("superseded", 111, 222))
    hidden.sheet_state = "hidden"
    book.save(str(_folder() / "Mixed.xlsx"))

    rec = _record(_list(), "Mixed.xlsx")
    text = (_folder() / rec["extracted_path"]).read_text(encoding="utf-8")
    marker = next(ln for ln in text.splitlines() if "Draft-Pricing" in ln)
    assert "[HIDDEN SHEET]" in marker, f"hidden sheet not marked: {marker!r}"
    # It is still extracted -- concealing it would be its own failure.
    assert "superseded" in text, "a hidden sheet's content should still be read"


def test_list_before_attach_errors():
    result = tt.cmd_list_attachments(SimpleNamespace(tender_id="NO-SUCH-TENDER"))
    assert "error" in result, "listing a nonexistent tender should error"


def main():
    setup_temp_vault()
    test_attach_creates_folder()
    test_attach_is_idempotent_and_keeps_provenance()
    test_attach_unknown_tender_rejected()
    test_manifest_written_with_lf()
    test_extracts_a_dropped_file()
    test_unsupported_type_still_listed()
    test_corrupt_file_is_distinguished_from_unsupported()
    test_no_text_layer_writes_no_file()
    test_unchanged_file_is_not_re_extracted()
    test_amended_document_detected_and_re_extracted()
    test_extractor_added_later_needs_no_re_drop()
    test_deleted_extracted_text_is_rebuilt()
    test_removed_file_drops_out_of_manifest()
    test_filename_with_awkward_characters_round_trips()
    test_read_is_paginated()
    test_read_offset_and_eof()
    test_read_limit_is_clamped()
    test_read_rejects_a_file_with_no_text()
    test_read_picks_up_a_change_without_a_list_call()
    test_same_stem_different_extension_keeps_separate_text()
    test_extracted_path_is_stable_across_amendments()
    test_unknown_platform_rejected()
    test_missing_library_is_reported_not_fatal()
    test_real_docx_extraction()
    test_docx_dropped_before_support_existed()
    test_real_xlsx_keeps_every_value_addressable()
    test_xlsx_uncalculated_formula_is_labelled_not_blank()
    test_xlsx_hidden_sheet_is_marked()
    test_list_before_attach_errors()
    print("All attachment tests passed.")


if __name__ == "__main__":
    main()

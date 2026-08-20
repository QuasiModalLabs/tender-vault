"""
Manually-dropped tender documents: extraction, manifest, drift detection.

THIS MODULE NEVER FETCHES ANYTHING. The RFP packages it reads are hosted on
Ariba and MERX behind account walls, and this project deliberately does not
scrape those platforms. A human pulls the package in a browser, past the
account wall, and drops the files into the tender's folder. Everything here
reads a directory that is already on disk. Adding a downloader to this file
would reverse a decision recorded in the README, not extend this one.

The extracted text stays out of git (`vault/tenders/` is ignored wholesale) and
out of the ChromaDB corpus. It is third-party content from a commercial
platform, and the extraction is that same content in another encoding.

WHAT GOES IN THE FOLDER

    vault/tenders/watching/cb-342-92719341/
        _index.md               this manifest
        RFP-W2187-SPO.pdf       dropped by hand
        .extracted/
            RFP-W2187-SPO.txt   what Claude actually reads

ADDING A FORMAT IS ONE FUNCTION AND ONE ROW

`EXTRACTORS` is a dispatch table keyed by lowercased file extension. No format
is special-cased anywhere in this module — `.pdf` is one row like any other.
Registering a new one means writing a function that returns an `ExtractResult`,
adding a row, and adding a line to requirements.txt. It needs no manifest schema
change, and it does not ask the user to re-drop files already in the folder: a
file sitting at `unsupported_type` is picked up on the next refresh once its
extension has an extractor (see `_needs_extraction`).

`.pdf`, `.docx` and `.xlsx` are registered, and they were added in that order
on purpose.

`.xlsx` was held back from the first two rounds, on the grounds that flattening
an evaluation grid or a pricing table into a stream of text loses the
row/column structure that gives every cell its meaning, producing confident
prose that is wrong about which figure belongs to which line item. That
objection was about FLATTENING, not about spreadsheets, and it is the reason
`_extract_xlsx` does not flatten: every value keeps its sheet, its column
letter and its real row number, so any figure can be cited as `Pricing!D7` and
checked against the file. A reader that cannot tell which row a number sits on
is the failure being avoided; a reader that can is doing what a person with the
file open does.

The risk moved rather than disappeared, so `_extract_xlsx` records what it
still cannot represent — number formats, merged-cell repeats, and the cached-
value trap that makes an uncalculated pricing column read as blank. Read that
docstring before trusting a number out of a spreadsheet.

A Word SOW is prose and never had this problem. It can still *contain* a
requirements table, which `_extract_docx` keeps one row per line rather than
dissolving into the surrounding text.

WHAT `page_count: null` IS FOR

Formats with no meaningful pagination report `null` rather than inventing a
number. That is what lets the manifest schema stay fixed as formats are added.

ONE LIMITATION, STATED OUT LOUD: upgrading an extractor does not by itself force
re-extraction. The refresh triggers on the file's hash and on whether an
extractor exists at all, not on the extractor's version — keeping the dispatch
table a plain function table was worth more than automatic version tracking. To
re-extract everything after an extractor upgrade, delete the folder's
`.extracted/` directory; the next refresh rebuilds it.
"""
from __future__ import annotations

import hashlib
import json
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Callable

import yaml


MANIFEST_NAME = "_index.md"
EXTRACTED_DIRNAME = ".extracted"

# The per-file enum. Three values, fixed: a file is readable, or an extractor
# looked and found no text, or nothing here can read it.
STATUS_EXTRACTED = "extracted"
STATUS_NO_TEXT_LAYER = "no_text_layer"
STATUS_UNSUPPORTED = "unsupported_type"

SOURCE_PLATFORMS = ("merx", "ariba", "other")

# Below this many non-whitespace characters, an extractor is reporting that the
# document has no text layer rather than that the document is short. Scaled per
# page with a floor, because "empty" for a 40-page scan is a much larger number
# than "empty" for a one-page cover letter.
_MIN_CHARS_PER_PAGE = 10
_MIN_CHARS_FLOOR = 50


@dataclass(frozen=True)
class ExtractResult:
    """What every extractor returns. `page_count` is None where meaningless."""
    text: str
    page_count: int | None
    extractor: str  # "name version", recorded in the manifest


# ---------------------------------------------------------------------------
# Extractors
# ---------------------------------------------------------------------------


def _extract_pdf(path: Path) -> ExtractResult:
    """
    pdfplumber. Imported lazily so this module loads without the dependency —
    a missing library is reported per-file and self-heals once it is installed,
    rather than taking down every attachment command with an ImportError.
    """
    import pdfplumber  # noqa: PLC0415 — lazy on purpose, see docstring

    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")

    version = getattr(pdfplumber, "__version__", "unknown")
    return ExtractResult(
        text="\n\n".join(pages),
        page_count=len(pages),
        extractor=f"pdfplumber {version}",
    )


def _extract_docx(path: Path) -> ExtractResult:
    """
    python-docx. Lazy import for the same reason as the PDF extractor.

    Walks the document body rather than reading `document.paragraphs` and
    `document.tables` separately, because those two lists lose the interleaving
    — and in a statement of work the table sits between the clauses that
    explain it, so reading all the prose and then all the tables reorders the
    document into something that never existed.

    Table rows are joined with ` | ` and kept one row per line. That is the
    limit of what flattening can honestly preserve: a requirements matrix stays
    row-shaped, which is enough to read a mandatory criterion against its
    response. It is also exactly why .xlsx is not here — see the module
    docstring.

    `page_count` is None. A .docx has no pages until something renders it, and
    the page count you see in Word is a property of the renderer plus the
    printer, not of the file. Reporting a number here would be inventing one.
    """
    import docx  # noqa: PLC0415 — lazy on purpose, see docstring
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = docx.Document(str(path))
    blocks: list[str] = []
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            blocks.append(Paragraph(child, document).text)
        elif child.tag.endswith("}tbl"):
            for row in Table(child, document).rows:
                blocks.append(" | ".join(cell.text.strip() for cell in row.cells))

    try:
        from importlib.metadata import version as _package_version
        installed = _package_version("python-docx")
    except Exception:
        installed = "unknown"

    return ExtractResult(
        text="\n".join(blocks),
        page_count=None,
        extractor=f"python-docx {installed}",
    )


_UNCALCULATED = "<uncalculated formula>"


def _xlsx_cell(value_cell, formula_cell) -> str:
    """
    One cell, with the formula trap made visible instead of silent.

    openpyxl with `data_only=True` returns the value Excel CACHED the last time
    it recalculated and saved. A workbook generated by a script, or saved by a
    tool that doesn't calculate, has no cache — every formula cell reads as
    None. On a pricing sheet that renders as a column of blanks where the
    numbers should be, which is worse than an error because it looks like the
    bidder left them empty. So a formula with no cached value is labelled, not
    blanked.
    """
    if value_cell.value is not None:
        return str(value_cell.value)
    raw = formula_cell.value
    if isinstance(raw, str) and raw.startswith("="):
        return _UNCALCULATED
    return ""


def _extract_xlsx(path: Path) -> ExtractResult:
    """
    openpyxl. Lazy import for the same reason as the other two.

    THE POINT OF THIS FORMAT IS THAT IT DOES NOT FLATTEN. The objection to
    reading spreadsheets was never that they are unreadable — it is that
    dissolving a grid into prose loses which figure belongs to which line item,
    so a reader confidently attributes a number to the wrong row. Every value
    here keeps its address: the sheet it came from, its column letter, and its
    row number, so any figure can be cited as `Pricing!D7` and checked. A
    reader that cannot tell which row a number is on is the failure; a reader
    that can is doing what it would do with the file open.

    Rows are emitted one per line, blank rows dropped, and row numbers are the
    real ones — so a gap in the numbering is a gap in the sheet.

    WHAT THIS STILL CANNOT DO, and it is recorded here because the risk moved
    rather than vanished:

      - Number formats are not applied. A cell holding 0.15 that Excel displays
        as "15%" extracts as 0.15, and a date may arrive as a datetime. Read a
        rate or a deadline off the source file, not off this text.
      - A merged cell holds its value only in the top-left of the range. The
        merged ranges are named in the sheet header so the shape is visible,
        but the repeat is not filled in.
      - Charts, images, comments and conditional formatting are not extracted.
      - Hidden sheets ARE extracted, and marked as hidden. A hidden pricing tab
        read as though it were the live one is its own kind of wrong.

    `page_count` is None. Sheets are not pages, and a spreadsheet has no
    pagination until something decides where to break it.
    """
    import openpyxl  # noqa: PLC0415 — lazy on purpose, see docstring

    values = openpyxl.load_workbook(path, data_only=True)
    formulas = openpyxl.load_workbook(path, data_only=False)

    blocks: list[str] = []
    uncalculated: list[str] = []
    total = len(values.worksheets)

    for index, sheet in enumerate(values.worksheets, start=1):
        formula_sheet = formulas[sheet.title]
        hidden = " [HIDDEN SHEET]" if sheet.sheet_state != "visible" else ""
        merged = [str(r) for r in sheet.merged_cells.ranges]
        header = (
            f"## Sheet {index} of {total}: {sheet.title!r}{hidden} — "
            f"{sheet.max_row} rows x {sheet.max_column} columns"
        )
        blocks.append(header)
        if merged:
            blocks.append(
                f"   merged ranges (value sits in the top-left only): "
                f"{', '.join(merged)}"
            )

        letters = [
            openpyxl.utils.get_column_letter(c)
            for c in range(1, sheet.max_column + 1)
        ]
        blocks.append("cols | " + " | ".join(letters))

        for row_index, (value_row, formula_row) in enumerate(
            zip(sheet.iter_rows(), formula_sheet.iter_rows()), start=1
        ):
            cells = [
                _xlsx_cell(v, f) for v, f in zip(value_row, formula_row)
            ]
            if not any(cell.strip() for cell in cells):
                continue                      # a blank row carries nothing
            for letter, cell in zip(letters, cells):
                if cell == _UNCALCULATED:
                    uncalculated.append(f"{sheet.title}!{letter}{row_index}")
            blocks.append(f"r{row_index} | " + " | ".join(cells))
        blocks.append("")

    if uncalculated:
        # At the top, where it is read before the numbers are.
        shown = ", ".join(uncalculated[:12])
        more = f" (+{len(uncalculated) - 12} more)" if len(uncalculated) > 12 else ""
        blocks.insert(0, "")
        blocks.insert(0, (
            f"!! {len(uncalculated)} formula cells have no cached value and are "
            f"shown as {_UNCALCULATED}: {shown}{more}. Excel stores results only "
            f"when it recalculates and saves, so these are unknown — NOT zero "
            f"and NOT blank."
        ))

    try:
        installed = getattr(openpyxl, "__version__", "unknown")
    except Exception:
        installed = "unknown"

    return ExtractResult(
        text="\n".join(blocks),
        page_count=None,
        extractor=f"openpyxl {installed}",
    )


# The dispatch table. One row per format; nothing below reads it by name.
EXTRACTORS: dict[str, Callable[[Path], ExtractResult]] = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".xlsx": _extract_xlsx,
}


# ---------------------------------------------------------------------------
# Hashing and manifest I/O
# ---------------------------------------------------------------------------


def sha256_file(path: Path) -> str:
    """Streamed, because an RFP package runs to tens of megabytes."""
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _scalar(value) -> str:
    """
    One YAML scalar, house style: hand-built, never yaml.dump.

    Strings go through json.dumps rather than the quote-swapping used for
    display titles elsewhere in the project (`title.replace('"', "'")`). That
    trick is fine for prose nobody looks up again; these strings include
    FILENAMES, which are keys — they have to come back byte-for-byte, and JSON
    string syntax is a subset of YAML 1.2 double-quoted syntax, so escaping is
    handled and round-trips exactly.

    `ensure_ascii=False` because this file is read in Obsidian: an accented
    filename or an em dash should appear as itself, not as `\\u2014`. The file
    is written UTF-8, which YAML requires anyway.
    """
    if value is None:
        return "null"
    if isinstance(value, bool):
        return "true" if value else "false"
    if isinstance(value, int):
        return str(value)
    return json.dumps(str(value), ensure_ascii=False)


# Field order in each file record. Fixed here so the manifest reads the same
# way every time it is rewritten and a diff shows only what actually changed.
_FILE_FIELDS = (
    "filename",
    "extension",
    "sha256",
    "size_bytes",
    "extraction_status",
    "extractor",
    "page_count",
    "char_count",
    "extracted_path",
    "first_seen",
    "last_changed",
    "superseded_sha256",
    "status_note",
)


def render_manifest(data: dict) -> str:
    """
    The manifest as markdown: authoritative YAML frontmatter, then a body for
    whoever opens the folder in Obsidian. The body is regenerated wholesale on
    every write; nothing reads it back.
    """
    lines = ["---"]
    for key in ("tender_id", "note", "source_platform", "retrieved_at",
                "manifest_updated_at"):
        if key in data:
            lines.append(f"{key}: {_scalar(data[key])}")

    lines.append("files:")
    for record in data.get("files", []):
        first = True
        for field in _FILE_FIELDS:
            if field not in record:
                continue
            prefix = "  - " if first else "    "
            lines.append(f"{prefix}{field}: {_scalar(record[field])}")
            first = False
        if first:  # a record with no recognised fields would break the sequence
            lines.append("  - {}")
    if not data.get("files"):
        lines[-1] = "files: []"
    lines.append("---")

    body = [
        "",
        f"# Documents — {data.get('tender_id', '')}",
        "",
        f"Tender note: {data.get('note', '')}",
        "",
        "Dropped by hand from "
        f"{data.get('source_platform', 'an unrecorded source')}. "
        "Nothing here was fetched by this project — see the module docstring in "
        "`scripts/attachments.py`.",
        "",
    ]
    files = data.get("files", [])
    if files:
        body += [
            "| File | Status | Pages | Extracted |",
            "| --- | --- | --- | --- |",
        ]
        for record in files:
            pages = record.get("page_count")
            body.append(
                f"| {record.get('filename', '')} "
                f"| {record.get('extraction_status', '')} "
                f"| {'—' if pages is None else pages} "
                f"| {record.get('extracted_path') or '—'} |"
            )
    else:
        body.append("_No documents dropped yet._")
    body.append("")
    return "\n".join(lines + body)


def write_manifest(folder: Path, data: dict) -> Path:
    """LF explicitly, as every markdown writer in this project does."""
    path = folder / MANIFEST_NAME
    path.write_text(render_manifest(data), encoding="utf-8", newline="\n")
    return path


def read_manifest(folder: Path) -> dict:
    """
    Parse the frontmatter with PyYAML, unlike the hand-rolled regex reader used
    for tender notes and digests.

    Those readers avoid safe_load for a specific reason (`_digest_frontmatter`
    in tender_tools): an unquoted `2026-08-09T14:27:11` loads as a datetime
    and then silently compares unequal to the string stamp it is checked
    against. That hazard is absent here because `_scalar` quotes every string it
    writes, so timestamps and hashes come back as `str`. And the regex reader
    could not do this job anyway — `files:` is a nested sequence of mappings,
    which `line.partition(":")` cannot represent.
    """
    path = folder / MANIFEST_NAME
    if not path.exists():
        return {}
    text = path.read_text(encoding="utf-8")
    match = re.match(r"^---\n(.*?)\n---", text, re.DOTALL)
    if not match:
        return {}
    loaded = yaml.safe_load(match.group(1)) or {}
    if not isinstance(loaded, dict):
        return {}
    loaded.setdefault("files", [])
    if not isinstance(loaded["files"], list):
        loaded["files"] = []
    return loaded


def new_manifest(tender_id: str, note_stem: str, source_platform: str) -> dict:
    return {
        "tender_id": tender_id,
        "note": f"[[{note_stem}]]",
        "source_platform": source_platform,
        "retrieved_at": datetime.now().strftime("%Y-%m-%d"),
        "manifest_updated_at": datetime.now().isoformat(timespec="seconds"),
        "files": [],
    }


# ---------------------------------------------------------------------------
# Refresh: diff the directory against the manifest
# ---------------------------------------------------------------------------


def dropped_files(folder: Path) -> list[Path]:
    """
    Everything a human put in the folder. Skips the manifest, `.extracted/`,
    and dotfiles — an OS writing `.DS_Store` into the folder is not a document.
    """
    out = []
    for child in sorted(folder.iterdir()):
        if child.is_dir() or child.name == MANIFEST_NAME:
            continue
        if child.name.startswith("."):
            continue
        out.append(child)
    return out


def _needs_extraction(record: dict | None, sha: str, ext: str, folder: Path) -> bool:
    """
    The whole re-extraction rule, in one place.

    Deliberately not a chain of special cases: a file is extracted again when
    what we recorded no longer describes what is on disk, or when this process
    can now read something it previously could not.
    """
    if record is None:
        return True                                  # never seen
    if record.get("sha256") != sha:
        return True                                  # amended document
    if record.get("extraction_status") == STATUS_UNSUPPORTED and ext in EXTRACTORS:
        return True                                  # an extractor arrived
    if record.get("extraction_status") == STATUS_EXTRACTED:
        rel = record.get("extracted_path")
        if not rel or not (folder / rel).exists():
            return True                              # text deleted by hand
    return False


def _classify(text: str, page_count: int | None) -> tuple[str, int]:
    """
    Distinguish 'the extractor found no text layer' from 'this document is
    short'. Returns (status, char_count); char_count goes into the manifest so
    the threshold's verdict is auditable instead of being a bare label.
    """
    chars = len(re.sub(r"\s", "", text))
    pages = page_count if page_count and page_count > 0 else 1
    threshold = max(_MIN_CHARS_FLOOR, _MIN_CHARS_PER_PAGE * pages)
    if chars < threshold:
        return STATUS_NO_TEXT_LAYER, chars
    return STATUS_EXTRACTED, chars


def _extracted_name(path: Path, claimed: set[str], reserved: dict[str, str]) -> str:
    """
    `RFP-W2187-SPO.pdf` -> `RFP-W2187-SPO.txt`, unless that name is taken.

    `reserved` maps a dropped filename to the .txt it already owns, and it is
    what makes this safe ACROSS runs rather than only within one. A file that
    is skipped as unchanged still owns its text, so without the reservation a
    newly-dropped `Annex-A.docx` would be handed `Annex-A.txt` — the name an
    untouched `Annex-A.pdf` was already using — and overwrite it. Both manifest
    records would then point at one file holding one document's text, which is
    the exact "confidently wrong" failure the hashes elsewhere exist to prevent.

    Returning a file's own reserved name also keeps the .txt path stable across
    re-extractions, so an amendment doesn't leave the old text orphaned.
    """
    own = reserved.get(path.name)
    if own:
        claimed.add(own)
        return own

    candidate = f"{path.stem}.txt"
    if candidate in claimed:
        candidate = f"{path.name}.txt"
    claimed.add(candidate)
    return candidate


def refresh(folder: Path) -> dict:
    """
    Diff the directory against the manifest, extract what changed, rewrite it.

    Returns the manifest plus what happened this call: `added`, `changed`
    (an amended document — the reason the hashes are here at all), `removed`,
    and `warnings` for conditions the per-file enum cannot express, such as a
    registered extractor whose library is not installed.
    """
    manifest = read_manifest(folder)
    if not manifest:
        manifest = {"files": []}

    previous = {r.get("filename"): r for r in manifest.get("files", [])}
    extracted_dir = folder / EXTRACTED_DIRNAME
    today = datetime.now().strftime("%Y-%m-%d")

    records: list[dict] = []
    added: list[str] = []
    changed: list[str] = []
    warnings: list[str] = []

    # Reserve the .txt names already owned by files still on disk, BEFORE
    # extracting anything. A file skipped as unchanged never reaches
    # _extracted_name, so its text would otherwise look unclaimed and be
    # handed to a different document. See _extracted_name.
    on_disk = dropped_files(folder)
    on_disk_names = {p.name for p in on_disk}
    claimed: set[str] = set()
    reserved: dict[str, str] = {}
    for old_record in manifest.get("files", []):
        name = old_record.get("filename")
        rel = old_record.get("extracted_path")
        if name in on_disk_names and rel:
            base = str(rel).rsplit("/", 1)[-1]
            reserved[name] = base
            claimed.add(base)

    for path in on_disk:
        ext = path.suffix.lower()
        sha = sha256_file(path)
        old = previous.get(path.name)

        record = {
            "filename": path.name,
            "extension": ext,
            "sha256": sha,
            "size_bytes": path.stat().st_size,
            "first_seen": (old or {}).get("first_seen", today),
            "last_changed": (old or {}).get("last_changed", today),
        }

        if old is None:
            added.append(path.name)
        elif old.get("sha256") != sha:
            # The case the hashing exists for. MERX posts addenda mid-
            # solicitation under the same filename, and extracted text that
            # looks current but describes the superseded document is worse than
            # having no text at all.
            changed.append(path.name)
            record["last_changed"] = today
            record["superseded_sha256"] = old.get("sha256")
        elif old.get("superseded_sha256"):
            record["superseded_sha256"] = old["superseded_sha256"]

        if not _needs_extraction(old, sha, ext, folder):
            for field in ("extraction_status", "extractor", "page_count",
                          "char_count", "extracted_path", "status_note"):
                if field in old:
                    record[field] = old[field]
            records.append(record)
            continue

        extractor = EXTRACTORS.get(ext)
        if extractor is None:
            # Listed, hashed and sized — visible even though unreadable. A
            # dropped pricing spreadsheet should not become invisible.
            record["extraction_status"] = STATUS_UNSUPPORTED
            record["extractor"] = None
            record["page_count"] = None
            record["extracted_path"] = None
            records.append(record)
            continue

        try:
            result = extractor(path)
        except ImportError as exc:
            # The extension is registered but its library is not installed.
            # Recorded as unsupported so `_needs_extraction` retries it once the
            # dependency appears, and surfaced as a warning so it does not read
            # as a fact about the document.
            warnings.append(
                f"{path.name}: no extractor available for {ext} in this "
                f"environment ({exc}). Install it and run this again."
            )
            record["extraction_status"] = STATUS_UNSUPPORTED
            record["extractor"] = None
            record["page_count"] = None
            record["extracted_path"] = None
            records.append(record)
            continue
        except Exception as exc:  # a malformed or encrypted document
            warnings.append(f"{path.name}: extraction failed ({exc}).")
            record["extraction_status"] = STATUS_UNSUPPORTED
            record["extractor"] = None
            record["page_count"] = None
            record["extracted_path"] = None
            records.append(record)
            continue

        status, char_count = _classify(result.text, result.page_count)
        record["extraction_status"] = status
        record["extractor"] = result.extractor
        record["page_count"] = result.page_count
        record["char_count"] = char_count

        if status == STATUS_NO_TEXT_LAYER:
            # No zero-byte .txt. A file that exists and is empty looks like a
            # successful extraction of an empty document, and the next reader
            # concludes the RFP said nothing.
            record["extracted_path"] = None
            # Deliberately not "this is a scan". That was true while .pdf was
            # the only registered format; a near-empty .docx reaches here too,
            # and telling the reader it is probably a scan would be a guess
            # stated as a finding.
            record["status_note"] = (
                "An extractor ran and found effectively no text. For a PDF that "
                "means a scan with no text layer. Not OCR'd — this is reported, "
                "not solved."
            )
        else:
            extracted_dir.mkdir(parents=True, exist_ok=True)
            name = _extracted_name(path, claimed, reserved)
            (extracted_dir / name).write_text(
                result.text, encoding="utf-8", newline="\n"
            )
            record["extracted_path"] = f"{EXTRACTED_DIRNAME}/{name}"

        records.append(record)

    removed = sorted(set(previous) - {r["filename"] for r in records})

    manifest["files"] = records
    manifest["manifest_updated_at"] = datetime.now().isoformat(timespec="seconds")
    write_manifest(folder, manifest)

    return {
        "manifest": manifest,
        "added": added,
        "changed": changed,
        "removed": removed,
        "warnings": warnings,
    }


def find_record(manifest: dict, filename: str) -> dict | None:
    for record in manifest.get("files", []):
        if record.get("filename") == filename:
            return record
    return None
